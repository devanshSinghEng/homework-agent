import os
from dotenv import load_dotenv
from pdf2image import convert_from_path
from google import genai
from docx import Document
import re
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
import sqlite3
from datetime import datetime

load_dotenv()
PDF_FILENAME = "pdf3.pdf"
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

def pdf_to_images(pdf_path):
    print("PDF se images banayi ja rahi hain...")
    images = convert_from_path(pdf_path, dpi=120)
    print(f"Total {len(images)} pages mile.")
    return images

def has_content(image, threshold=0.02):
    gray = image.convert("L")
    arr = np.array(gray)
    bright_pixels = np.sum(arr > 60)
    total_pixels = arr.size
    ratio = bright_pixels / total_pixels
    return ratio > threshold

def setup_database():
    conn = sqlite3.connect("homework_tracker.db")
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS homework_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            source_pdf TEXT,
            topic TEXT,
            num_questions INTEGER,
            generated_on TEXT
        )
    """)
    conn.commit()
    conn.close()

def add_formatted_text(doc, text):
    lines = text.split("\n")
    for line in lines:
        line = line.strip()

        if line == "" or line == "---":
            continue

        line = line.replace("`", "")

        if line.startswith(">>> "):
            code_para = doc.add_paragraph()
            code_para.paragraph_format.left_indent = Inches(0.4)
            run = code_para.add_run(line.replace(">>> ", ""))
            run.font.name = "Consolas"
            run.font.size = Pt(11)
            continue

        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(10)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)

setup_database()

pages = pdf_to_images(PDF_FILENAME)
pages[0].save("test_page.png")

pages_with_content = [p for p in pages if has_content(p)]
print(f"Total {len(pages)} pages the, jinme se {len(pages_with_content)} pages mein actual content mila.")

prompt = """These are a student's handwritten class notes (multiple pages).
Carefully read ALL the pages and create ONE combined homework sheet covering the topics from the entire notes.

First, output the topic name on its own line in this exact format:
TOPIC: <topic name here>

Then, after a blank line, provide ONLY:
1. A list of 10 practice questions (concept-based, easy to medium difficulty)

Do NOT include answers.
Do NOT add any title or heading like "Homework Sheet" - start directly with the questions after the topic line.
Do NOT use markdown code formatting (no backticks).

IMPORTANT: If a question refers to a code statement (like a variable declaration),
write the question first, then put the code statement on its own new line below it,
starting that line with ">>> " (greater-than symbols followed by a space).
Example:
5. Identify the data type, variable name, and value in the statement below.
>>> int age = 25;

Write everything in clear, simple English only. Do not use Hindi or Hinglish anywhere.
Only give the topic line and the numbered questions, no extra commentary."""

print(f"{len(pages_with_content)} pages Gemini ko bheji ja rahi hain, thoda time lagega...")

response = client.models.generate_content(
    model="gemini-3.6-flash",
    contents=[prompt] + pages_with_content
)

print("\n--- AI se mila homework ---\n")
print(response.text)
response_lines = response.text.strip().split("\n")
topic = "Unknown"
remaining_lines = []

for line in response_lines:
    if line.strip().startswith("TOPIC:"):
        topic = line.replace("TOPIC:", "").strip()
    else:
        remaining_lines.append(line)

questions_text = "\n".join(remaining_lines)
num_questions = len([l for l in remaining_lines if re.match(r'^\d+\.', l.strip())])

print(f"\nTopic mila: {topic}")
print(f"Total questions: {num_questions}")

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(14)
heading = doc.add_heading("Homework Sheet", level=0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_text(doc, questions_text)
doc.save("homework_output.docx")

conn = sqlite3.connect("homework_tracker.db")
cursor = conn.cursor()
cursor.execute("""
    INSERT INTO homework_log (source_pdf, topic, num_questions, generated_on)
    VALUES (?, ?, ?, ?)
""", (PDF_FILENAME, topic, num_questions, datetime.now().strftime("%Y-%m-%d %H:%M")))
conn.commit()
conn.close()

print("Database mein bhi record save ho gaya!")

print("\nHomework sheet 'homework_output.docx' mein save ho gayi!")